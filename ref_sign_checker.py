from docx import Document
from docx.shared import RGBColor
from docx.enum.section import WD_ORIENT
import unicodedata
import re
from collections import Counter
import os

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ========== 1. Constants / Patterns ==========

EN_NUM_MAP = {
    "1st": "1", "2nd": "2", "3rd": "3", "4th": "4", "5th": "5",
    "6th": "6", "7th": "7", "8th": "8", "9th": "9",
    "first": "1", "second": "2", "third": "3", "fourth": "4", "fifth": "5",
    "sixth": "6", "seventh": "7", "eighth": "8", "ninth": "9",
    "one": "1", "two": "2", "three": "3", "four": "4", "five": "5",
    "six": "6", "seven": "7", "eight": "8", "nine": "9", "ten": "10", "zero": "0",
}

# Common units in patents (for "12 mg" → "12mg" normalization)
UNITS = {
    # Mass
    "mg", "g", "kg", "µg", "ug", "ng", "pg", "t",
    # Length
    "mm", "cm", "m", "km", "µm", "um", "nm", "pm", "fm", "Å",
    # Area
    "mm2", "cm2", "m2", "km2",
    # Volume
    "ml", "mL", "l", "L", "µl", "µL", "uL", "nL", "cc",
    "mm3", "cm3", "m3",
    # Time
    "ms", "µs", "us", "ns", "ps", "fs", "s", "sec", "min", "h", "hr", "hrs",
    # Electrical
    "mV", "V", "kV", "MV",
    "mA", "A", "kA", "µA", "uA", "nA", "pA",
    "mW", "W", "kW", "MW", "GW",
    "Ω", "ohm", "kΩ", "MΩ", "mΩ",
    "F", "mF", "µF", "uF", "nF", "pF",
    "H", "mH", "µH", "uH", "nH",
    "C", "mC", "µC", "nC",
    # Frequency
    "Hz", "kHz", "MHz", "GHz", "THz",
    # Pressure
    "Pa", "kPa", "MPa", "GPa", "hPa", "bar", "mbar", "atm", "torr", "Torr",
    # Temperature
    "K", "°C", "°F", "degC", "degF",
    # Concentration / Chemistry
    "mol", "mmol", "µmol", "umol", "nmol", "pmol",
    "M", "mM", "µM", "uM", "nM", "pM",
    "ppm", "ppb", "ppt",
    "wt", "vol",
    # Data
    "bit", "B", "KB", "MB", "GB", "TB", "PB", "kB",
    "bps", "Kbps", "Mbps", "Gbps",
    # Angle
    "rad", "mrad", "deg", "°",
    # Force / Energy
    "N", "kN", "mN", "µN",
    "J", "kJ", "mJ", "µJ", "MJ", "GJ",
    "eV", "keV", "MeV", "GeV",
    "cal", "kcal",
    "Wh", "kWh", "MWh",
    # Power density, etc.
    "W/m2", "mW/cm2",
    # Speed
    "m/s", "km/h", "mph", "rpm", "rps",
    # Misc
    "dB", "dBm", "dBi",
    "lm", "lx", "cd",
    "Bq", "Sv", "Gy", "rem", "rad",
    "mol/L", "g/L", "mg/L", "µg/L",
    "g/mol", "kg/mol",
    "N/m", "N/mm", "kN/m",
    "Pa·s", "mPa·s",
    "S", "mS", "µS",  # Siemens
    "T", "mT", "µT", "G",  # Tesla, Gauss
    "Wb",  # Weber
}

# Build pattern for number + space + unit (sorted by length desc to match longer units first)
# Exclude single-letter units (A, V, W, etc.) to avoid matching reference signs like "31 A"
_units_multi_char = {u for u in UNITS if len(u) >= 2}
_units_pattern = "|".join(re.escape(u) for u in sorted(_units_multi_char, key=len, reverse=True))
pat_en_num_space_unit = re.compile(
    r"(\d+(?:[.,]\d+)?)\s+(" + _units_pattern + r")(?![A-Za-z])"
)

# EN patterns
pat_en_1 = re.compile(r"[A-Za-z0-9Α-Ωα-ω]*[0-9Α-Ωα-ω]+[A-Za-z0-9Α-Ωα-ω]*")
pat_en_2 = re.compile(r"\b[A-Z_]{2,}\b|(?:(?<=\W)|^)[A-Z_](?=\W)")
pat_en_num_commas = re.compile(r"\d{1,3}(?:,\d{3})+")
pat_en_caps_plural = re.compile(r"\b([A-Z_]{2,})s\b")

# ALLCAPS + space(s) + digits (e.g., "UDM 12")
pat_en_caps_num_space = re.compile(r"\b([A-Z_]{2,})\s+(\d+)\b")

# JP patterns
pat_jp_num = re.compile(r"\d+")
pat_jp_alnum = re.compile(r"[A-Za-z0-9Α-Ωα-ω]*[0-9Α-Ωα-ω]+[A-Za-z0-9Α-Ωα-ω]*")
pat_jp_alpha = re.compile(r"[A-Za-zΑ-Ωα-ω_]+")
pat_jp_zero = re.compile(r"ゼロ")


# ========== 2. Utilities ==========

def nfkc(s: str) -> str:
    return unicodedata.normalize("NFKC", s or "")

def strip_num_commas(s: str) -> str:
    return (s or "").replace(",", "")

def first_alpha_match(text: str):
    m = re.search(r"[A-Za-z]+", text)
    return (m.group(0), m.span()) if m else (None, None)


# ========== TL direction detection ==========

def is_japanese_char(ch: str) -> bool:
    code = ord(ch)
    if 0x3040 <= code <= 0x309F:
        return True  # Hiragana
    if 0x30A0 <= code <= 0x30FF:
        return True  # Katakana
    if 0x4E00 <= code <= 0x9FFF:
        return True  # Kanji
    if 0xFF00 <= code <= 0xFFEF:
        return True  # Full-width forms
    return False

def jp_score(text: str) -> float:
    chars = [c for c in text if not c.isspace()]
    if not chars:
        return 0.0
    jp = sum(1 for c in chars if is_japanese_char(c))
    return jp / len(chars)

def en_score(text: str) -> float:
    chars = [c for c in text if not c.isspace()]
    if not chars:
        return 0.0
    en = sum(1 for c in chars if ("A" <= c <= "Z") or ("a" <= c <= "z"))
    return en / len(chars)

def detect_column_lang(doc: Document, max_rows=50):
    for tbl in doc.tables:
        col_scores = [{"jp": 0.0, "en": 0.0}, {"jp": 0.0, "en": 0.0}]
        checked = 0

        for row in tbl.rows:
            if len(row.cells) < 2:
                continue

            t0 = nfkc(row.cells[0].text)
            t1 = nfkc(row.cells[1].text)

            col_scores[0]["jp"] += jp_score(t0)
            col_scores[0]["en"] += en_score(t0)
            col_scores[1]["jp"] += jp_score(t1)
            col_scores[1]["en"] += en_score(t1)

            checked += 1
            if checked >= max_rows:
                break

        if checked > 0:
            jp_col = 0 if col_scores[0]["jp"] >= col_scores[1]["jp"] else 1
            en_col = 1 - jp_col
            return jp_col, en_col

    return 0, 1  # fallback

def direction_tag(jp_col, en_col):
    return "JP2EN" if jp_col < en_col else "EN2JP"


# ========== 3. JP Tokenizer ==========

def tokenize_jp(text: str):
    text = nfkc(text)
    tokens = []
    seen = set()

    # ゼロ → 0
    for m in pat_jp_zero.finditer(text):
        span = (m.start(), m.end())
        if span in seen:
            continue
        seen.add(span)
        tokens.append({
            "surface": m.group(0),
            "norm": "0",
            "start": m.start(),
            "end": m.end(),
            "cls": "jp_number",
        })

    # numbers with commas
    for m in pat_en_num_commas.finditer(text):
        span = (m.start(), m.end())
        if span in seen:
            continue
        seen.add(span)
        tok = m.group(0)
        tokens.append({
            "surface": tok,
            "norm": strip_num_commas(tok),
            "start": m.start(),
            "end": m.end(),
            "cls": "jp_number",
        })

    # plain numbers
    for m in pat_jp_num.finditer(text):
        span = (m.start(), m.end())
        if span in seen:
            continue
        seen.add(span)
        tok = m.group(0)
        tokens.append({
            "surface": tok,
            "norm": strip_num_commas(tok),
            "start": m.start(),
            "end": m.end(),
            "cls": "jp_number",
        })

    # alnum with digits (S10, RAM100, 5A, UDM12, 12mg, etc.)
    for m in pat_jp_alnum.finditer(text):
        span = (m.start(), m.end())
        if span in seen:
            continue
        seen.add(span)
        tok = m.group(0)
        tokens.append({
            "surface": tok,
            "norm": strip_num_commas(tok),
            "start": m.start(),
            "end": m.end(),
            "cls": "jp_alnum_digit",
        })

    alnum_spans = [(t["start"], t["end"]) for t in tokens if t["cls"] == "jp_alnum_digit"]

    def is_inside_any_alnum(span):
        s, e = span
        return any(a_s <= s and e <= a_e for a_s, a_e in alnum_spans)

    # roman alpha in JP text
    for m in pat_jp_alpha.finditer(text):
        span = (m.start(), m.end())
        if is_inside_any_alnum(span):
            continue
        if span in seen:
            continue
        seen.add(span)
        tok = m.group(0)
        tokens.append({
            "surface": tok,
            "norm": tok,
            "start": m.start(),
            "end": m.end(),
            "cls": "jp_alpha",
        })

    return tokens

def mark_embedded_numbers(jp_tokens):
    alnum_spans = [(t["start"], t["end"]) for t in jp_tokens if t["cls"] == "jp_alnum_digit"]
    for t in jp_tokens:
        t["embedded"] = False
        if t["cls"] == "jp_number":
            s, e = t["start"], t["end"]
            for a_s, a_e in alnum_spans:
                if a_s <= s and e <= a_e:
                    t["embedded"] = True
                    break


# ========== 4. EN Tokenizer ==========

def tokenize_en(text: str):
    text = nfkc(text)
    tokens = []
    seen = set()

    first_alpha, first_span = first_alpha_match(text)

    # Number + space + unit → normalize to "numberunit" (e.g., "12 mg" → "12mg")
    # This handles JP "１２ｍｇ" (→ "12mg" after NFKC) vs EN "12 mg"
    for m in pat_en_num_space_unit.finditer(text):
        span = (m.start(), m.end())
        if span in seen:
            continue
        seen.add(span)
        num, unit = m.group(1), m.group(2)
        # Normalize: just remove space, keep original case
        tokens.append({
            "surface": m.group(0),
            "norm": f"{num}{unit}",
            "start": m.start(),
            "end": m.end(),
            "cls": "en_num_unit",
        })

    # plural ALLCAPS like APIs → normalize to API
    for m in pat_en_caps_plural.finditer(text):
        span = (m.start(), m.end())
        if span in seen:
            continue
        seen.add(span)
        base = m.group(1)
        surface = m.group(0)
        tokens.append({
            "surface": surface,
            "norm": base,
            "start": m.start(),
            "end": m.end(),
            "cls": "en_caps_plural",
        })

    # ALLCAPS + space(s) + digits combo candidate (e.g., 'UDM 12' -> 'UDM12')
    for m in pat_en_caps_num_space.finditer(text):
        full_span = (m.start(), m.end())
        if full_span in seen:
            continue
        alpha = m.group(1)
        num = m.group(2)

        seen.add(full_span)
        seen.add((m.start(1), m.end(1)))
        seen.add((m.start(2), m.end(2)))

        tokens.append({
            "surface": m.group(0),
            "norm": f"{alpha}{num}",
            "start": m.start(),
            "end": m.end(),
            "cls": "en_caps_num_combo",
            "alpha": alpha,
            "num": num,
            "alpha_span": (m.start(1), m.end(1)),
            "num_span": (m.start(2), m.end(2)),
        })

    # numbers with commas
    for m in pat_en_num_commas.finditer(text):
        span = (m.start(), m.end())
        if span in seen:
            continue
        seen.add(span)
        tok = m.group(0)
        tokens.append({
            "surface": tok,
            "norm": strip_num_commas(tok),
            "start": m.start(),
            "end": m.end(),
            "cls": "en_alnum_digit",
        })

    # alnum with digits
    for m in pat_en_1.finditer(text):
        span = (m.start(), m.end())
        if span in seen:
            continue
        seen.add(span)
        tok = m.group(0)
        tokens.append({
            "surface": tok,
            "norm": strip_num_commas(tok),
            "start": m.start(),
            "end": m.end(),
            "cls": "en_alnum_digit",
        })

    # caps tokens (including single-letter caps, except leading article "A" case)
    for m in pat_en_2.finditer(text):
        span = (m.start(), m.end())
        tok = m.group(0)
        if tok == "A" and first_alpha == "A" and first_span == span:
            continue
        if span in seen:
            continue
        seen.add(span)
        tokens.append({
            "surface": tok,
            "norm": tok,
            "start": m.start(),
            "end": m.end(),
            "cls": "en_caps",
        })

    # number words
    word_pat = re.compile(r"\b[a-zA-Z]+\b")
    for m in word_pat.finditer(text):
        w = m.group(0).lower()
        if w in EN_NUM_MAP:
            span = (m.start(), m.end())
            if span in seen:
                continue
            seen.add(span)
            tokens.append({
                "surface": m.group(0),
                "norm": EN_NUM_MAP[w],
                "start": m.start(),
                "end": m.end(),
                "cls": "en_num_word",
            })

    return tokens

def apply_conditional_en_caps_num_combo(jp_tokens, en_tokens):
    """
    If EN has a combo token like 'UDM 12' -> norm 'UDM12', keep it ONLY when JP contains 'UDM12'.
    Otherwise decompose it back into 'UDM' and '12'.
    """
    jp_norms = set(t["norm"] for t in jp_tokens)

    out = []
    for t in en_tokens:
        if t.get("cls") != "en_caps_num_combo":
            out.append(t)
            continue

        combo_norm = t.get("norm", "")
        if combo_norm and combo_norm in jp_norms:
            out.append(t)
        else:
            a_s, a_e = t["alpha_span"]
            n_s, n_e = t["num_span"]
            alpha = t["alpha"]
            num = t["num"]

            out.append({
                "surface": alpha,
                "norm": alpha,
                "start": a_s,
                "end": a_e,
                "cls": "en_caps",
            })
            out.append({
                "surface": num,
                "norm": num,
                "start": n_s,
                "end": n_e,
                "cls": "en_alnum_digit",
            })

    return out

def counter_from_tokens(tokens):
    return Counter(t["norm"] for t in tokens)


# ========== 5. Alpha cross-match ==========

def find_alpha_in_en(alpha, en_text):
    pat = r"(?<![A-Za-z0-9])" + re.escape(alpha) + r"(?![A-Za-z0-9])"
    return re.search(pat, en_text, flags=re.I)

def cross_match_jp_alpha(jp_tokens, en_text):
    jp_shared, jp_only, en_shared = [], [], []

    for t in jp_tokens:
        if t["cls"] != "jp_alpha":
            continue
        tok = t["surface"]
        if not tok:
            continue
        if tok.isupper():
            continue  # CAPS handled elsewhere

        m = find_alpha_in_en(tok, en_text)
        if m:
            jp_shared.append((t["start"], t["end"]))
            en_shared.append((m.start(), m.end()))
        else:
            jp_only.append((t["start"], t["end"]))

    return jp_shared, jp_only, en_shared


# ========== 6. Evidence-based split ==========

def evidence_split_jp_alnum(jp_tokens, en_tokens, en_text, shared_spans_jp, only_spans_jp, shared_spans_en):
    """
    Handle compound reference signs where JP has concatenated form and EN has split form.
    
    Case B (ABC12 style): e.g., JP "RAM100" vs EN "RAM 100" (handled via ALLCAPS+num combo)
    
    Note: Case A (12ABC style like "32A") was removed because it caused false matches
    when single letters like "A" matched English articles ("a", "an").
    """
    en_norms = set(t["norm"] for t in en_tokens)
    to_remove = set()

    def add_shared(sp):
        if sp not in shared_spans_jp:
            shared_spans_jp.append(sp)

    for t in jp_tokens:
        if t["cls"] != "jp_alnum_digit":
            continue

        s_full = t["surface"]
        base = t["start"]
        full_sp = (t["start"], t["end"])

        # Case B: ABC12 style (e.g., "RAM100", "Device50")
        # Only match if BOTH alpha and num are found as separate tokens in EN
        mB = re.match(r"^([A-Za-zΑ-Ωα-ω]+)(\d+)$", s_full)
        if mB:
            alpha, num = mB.group(1), mB.group(2)
            if alpha in en_norms and num in en_norms:
                alpha_sp = (base, base + len(alpha))
                num_sp = (base + len(alpha), t["end"])
                add_shared(alpha_sp)
                add_shared(num_sp)
                to_remove.update([full_sp, alpha_sp, num_sp])
            continue

    if to_remove:
        only_spans_jp[:] = [sp for sp in only_spans_jp if sp not in to_remove]


# ========== 7. Span helpers for mismatch logic ==========

def spans_contain(outer, inner):
    os_, oe = outer
    is_, ie = inner
    return os_ <= is_ and ie <= oe

def filter_inner_only(only_spans, shared_spans):
    if not only_spans:
        return []
    if not shared_spans:
        return only_spans
    out = []
    for sp in only_spans:
        if any(spans_contain(sh, sp) for sh in shared_spans):
            continue
        out.append(sp)
    return out


# ========== 8. Dark table styling + Highlighting ==========

WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK_HEX = "000000"
WHITE_HEX = "FFFFFF"
RED = RGBColor(0xFF, 0x00, 0x00)
YELLOW = RGBColor(0xFF, 0xFF, 0x00)

def _set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)

def _set_cell_borders(cell, color_hex: str = WHITE_HEX, size: str = "8"):
    """
    size is in eighths of a point (8 = 1pt). Use 4 for ~0.5pt if you prefer.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color_hex)

def _set_runs_white(cell):
    """
    Make all existing runs white (used for cells not rewritten by mismatch highlighter).
    Preserves RED (mismatch) and YELLOW (match) highlighting.
    """
    for p in cell.paragraphs:
        for r in p.runs:
            if r.font.color and r.font.color.rgb in (RED, YELLOW):
                continue
            r.font.color.rgb = WHITE

def apply_highlighting_to_cell(cell, shared_spans, mismatch_spans, default_color=WHITE):
    """
    Rewrite cell's first paragraph into runs:
      - shared spans: yellow + bold
      - mismatch spans: red + bold
      - other text: default_color (white)
    
    Mismatch (red) takes priority over shared (yellow) if spans overlap.
    """
    if not shared_spans and not mismatch_spans:
        _set_runs_white(cell)
        return

    text = cell.text
    if not text:
        return

    para = cell.paragraphs[0]

    # clear existing runs
    for r in list(para.runs):
        r.clear()
        r.text = ""

    # Build a list of (start, end, color, bold) segments
    # Start with all text as default
    # Then apply shared (yellow), then mismatch (red) to override
    
    # Create a character-level color map
    char_colors = [default_color] * len(text)
    char_bold = [False] * len(text)
    
    # Apply shared spans (yellow) first
    for s, e in shared_spans:
        for i in range(s, min(e, len(text))):
            char_colors[i] = YELLOW
            char_bold[i] = True
    
    # Apply mismatch spans (red) - overrides yellow
    for s, e in mismatch_spans:
        for i in range(s, min(e, len(text))):
            char_colors[i] = RED
            char_bold[i] = True
    
    # Convert character-level map to runs (consecutive chars with same formatting)
    if not text:
        return
        
    pos = 0
    while pos < len(text):
        # Find the end of this run (consecutive chars with same color/bold)
        run_color = char_colors[pos]
        run_bold = char_bold[pos]
        end = pos + 1
        while end < len(text) and char_colors[end] == run_color and char_bold[end] == run_bold:
            end += 1
        
        run = para.add_run(text[pos:end])
        run.font.color.rgb = run_color
        run.font.bold = run_bold
        pos = end


def apply_red_bold_to_span(cell, spans, default_color=WHITE):
    """
    Legacy function for backward compatibility.
    Rewrite cell's first paragraph into runs:
      - normal runs: white
      - mismatch spans: red + bold
    """
    if not spans:
        _set_runs_white(cell)
        return

    text = cell.text
    if not text:
        return

    para = cell.paragraphs[0]

    # clear existing runs
    for r in list(para.runs):
        r.clear()
        r.text = ""

    spans_sorted = sorted(spans, key=lambda x: (x[0], x[1]))
    merged = []
    for s, e in spans_sorted:
        if not merged:
            merged.append([s, e])
            continue
        ls, le = merged[-1]
        if s <= le:
            merged[-1][1] = max(le, e)
        else:
            merged.append([s, e])

    pos = 0
    for s, e in merged:
        if pos < s:
            run = para.add_run(text[pos:s])
            run.font.color.rgb = default_color
        run = para.add_run(text[s:e])
        run.font.bold = True
        run.font.color.rgb = RED
        pos = e

    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.color.rgb = default_color

def set_document_landscape(doc: Document) -> None:
    """Set all sections to landscape."""
    for section in doc.sections:
        if section.orientation != WD_ORIENT.LANDSCAPE:
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height

def apply_dark_table_style(tbl):
    """
    Black background, white grid lines, white base text.
    (Mismatch red/bold remains as-is.)
    """
    for row in tbl.rows:
        for cell in row.cells:
            _set_cell_shading(cell, BLACK_HEX)
            _set_cell_borders(cell, WHITE_HEX, size="8")
            _set_runs_white(cell)


def distribute_column_widths(tbl):
    """
    Set fixed column widths: Seg=13mm, Source=122mm, Target=122mm.
    This mimics Word's「列の幅を揃える」for the 2nd and 3rd columns.
    """
    from docx.shared import Mm
    
    seg_width = Mm(13)
    content_col_width = Mm(122)
    
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 3:
            cells[0].width = seg_width
            cells[1].width = content_col_width
            cells[2].width = content_col_width


# ========== 9. Segment numbering column ==========

def table_has_trados_header(tbl) -> bool:
    """
    Trados: header row exists, and first cell is 'Source segment' (exact match).
    Phrase: usually no header row.
    """
    try:
        v = nfkc(tbl.rows[0].cells[0].text).strip()
    except Exception:
        return False
    return v == "Source segment"

def insert_left_cell(row):
    """
    Insert a new cell at the left of a row (index 0).
    Returns the newly inserted cell.
    """
    new_tc = OxmlElement("w:tc")
    row._tr.insert(0, new_tc)
    return row.cells[0]

def add_segment_number_column(tbl, has_header: bool, start_at=1, header_label="Seg"):
    """
    Adds a narrow left column with segment numbers.

    - If has_header=True:
        row 0 gets header_label, numbering starts from row 1 as 1..N
    - If has_header=False:
        numbering starts from row 0 as 1..N (no header row)
    """
    for r_idx, row in enumerate(tbl.rows):
        c = insert_left_cell(row)

        if has_header and r_idx == 0:
            c.text = str(header_label)
        else:
            offset = 1 if has_header else 0
            c.text = str(start_at + (r_idx - offset))

        # Center + bold
        if c.paragraphs:
            c.paragraphs[0].alignment = 1  # center
        for p in c.paragraphs:
            for run in p.runs:
                run.font.bold = True


# ========== 10. Main processing ==========

def process_docx(in_path, out_path_mismatch, out_path_all):
    """
    Process the DOCX and output two files:
    - out_path_mismatch: only rows with mismatches (red highlighting)
    - out_path_all: all rows with both matches (yellow) and mismatches (red)
    """
    # We need to process twice: once for all-rows, once for mismatch-only
    # Process all-rows first, then mismatch-only
    
    # --- First pass: All rows output ---
    doc_all = Document(in_path)
    set_document_landscape(doc_all)
    jp_col, en_col = detect_column_lang(doc_all, max_rows=50)

    for tbl in doc_all.tables:
        has_header = table_has_trados_header(tbl)
        add_segment_number_column(tbl, has_header=has_header, start_at=1, header_label="Seg")

        jp_col_shifted = jp_col + 1
        en_col_shifted = en_col + 1

        for row_idx, row in enumerate(tbl.rows):
            if has_header and row_idx == 0:
                continue
            if len(row.cells) < 3:
                continue

            jp_cell = row.cells[jp_col_shifted]
            en_cell = row.cells[en_col_shifted]

            jp_text = nfkc(jp_cell.text)
            en_text = nfkc(en_cell.text)

            jp_t = tokenize_jp(jp_text)
            en_t = tokenize_en(en_text)
            mark_embedded_numbers(jp_t)
            en_t = apply_conditional_en_caps_num_combo(jp_t, en_t)

            jp_cnt = counter_from_tokens(jp_t)
            en_cnt = counter_from_tokens(en_t)

            shared = jp_cnt & en_cnt
            jp_only = jp_cnt - shared
            en_only = en_cnt - shared

            # Build shared spans
            shared_jp, only_jp = [], []
            tmp_s = shared.copy()
            tmp_jp = jp_only.copy()

            for t in jp_t:
                if t.get("embedded"):
                    continue
                n = t["norm"]
                if tmp_s.get(n, 0) > 0:
                    shared_jp.append((t["start"], t["end"]))
                    tmp_s[n] -= 1

            for t in jp_t:
                if not t.get("embedded"):
                    continue
                n = t["norm"]
                if tmp_s.get(n, 0) > 0:
                    shared_jp.append((t["start"], t["end"]))
                    tmp_s[n] -= 1

            for t in jp_t:
                n = t["norm"]
                if tmp_jp.get(n, 0) > 0:
                    only_jp.append((t["start"], t["end"]))
                    tmp_jp[n] -= 1

            tmp_s2 = shared.copy()
            tmp_en = en_only.copy()
            shared_en, only_en = [], []

            for t in en_t:
                n = t["norm"]
                if tmp_s2.get(n, 0) > 0:
                    shared_en.append((t["start"], t["end"]))
                    tmp_s2[n] -= 1
                elif tmp_en.get(n, 0) > 0:
                    only_en.append((t["start"], t["end"]))
                    tmp_en[n] -= 1

            jp_alpha_s, jp_alpha_o, en_alpha_s = cross_match_jp_alpha(jp_t, en_text)
            for sp in jp_alpha_s:
                if sp in only_jp:
                    only_jp.remove(sp)
            shared_jp.extend(jp_alpha_s)
            shared_en.extend(en_alpha_s)
            only_jp.extend(jp_alpha_o)

            evidence_split_jp_alnum(jp_t, en_t, en_text, shared_jp, only_jp, shared_en)

            only_jp = filter_inner_only(only_jp, shared_jp)
            only_en = filter_inner_only(only_en, shared_en)

            # Apply highlighting: yellow for shared, red for mismatch
            apply_highlighting_to_cell(jp_cell, shared_jp, only_jp, default_color=WHITE)
            apply_highlighting_to_cell(en_cell, shared_en, only_en, default_color=WHITE)

        apply_dark_table_style(tbl)
        distribute_column_widths(tbl)

    doc_all.save(out_path_all)

    # --- Second pass: Mismatch-only output ---
    doc_mismatch = Document(in_path)
    set_document_landscape(doc_mismatch)

    for tbl in doc_mismatch.tables:
        has_header = table_has_trados_header(tbl)
        add_segment_number_column(tbl, has_header=has_header, start_at=1, header_label="Seg")

        jp_col_shifted = jp_col + 1
        en_col_shifted = en_col + 1

        rows_to_delete_idx = []

        for row_idx, row in enumerate(tbl.rows):
            if has_header and row_idx == 0:
                continue
            if len(row.cells) < 3:
                continue

            jp_cell = row.cells[jp_col_shifted]
            en_cell = row.cells[en_col_shifted]

            jp_text = nfkc(jp_cell.text)
            en_text = nfkc(en_cell.text)

            jp_t = tokenize_jp(jp_text)
            en_t = tokenize_en(en_text)
            mark_embedded_numbers(jp_t)
            en_t = apply_conditional_en_caps_num_combo(jp_t, en_t)

            jp_cnt = counter_from_tokens(jp_t)
            en_cnt = counter_from_tokens(en_t)

            shared = jp_cnt & en_cnt
            jp_only = jp_cnt - shared
            en_only = en_cnt - shared

            shared_jp, only_jp = [], []
            tmp_s = shared.copy()
            tmp_jp = jp_only.copy()

            for t in jp_t:
                if t.get("embedded"):
                    continue
                n = t["norm"]
                if tmp_s.get(n, 0) > 0:
                    shared_jp.append((t["start"], t["end"]))
                    tmp_s[n] -= 1

            for t in jp_t:
                if not t.get("embedded"):
                    continue
                n = t["norm"]
                if tmp_s.get(n, 0) > 0:
                    shared_jp.append((t["start"], t["end"]))
                    tmp_s[n] -= 1

            for t in jp_t:
                n = t["norm"]
                if tmp_jp.get(n, 0) > 0:
                    only_jp.append((t["start"], t["end"]))
                    tmp_jp[n] -= 1

            tmp_s2 = shared.copy()
            tmp_en = en_only.copy()
            shared_en, only_en = [], []

            for t in en_t:
                n = t["norm"]
                if tmp_s2.get(n, 0) > 0:
                    shared_en.append((t["start"], t["end"]))
                    tmp_s2[n] -= 1
                elif tmp_en.get(n, 0) > 0:
                    only_en.append((t["start"], t["end"]))
                    tmp_en[n] -= 1

            jp_alpha_s, jp_alpha_o, en_alpha_s = cross_match_jp_alpha(jp_t, en_text)
            for sp in jp_alpha_s:
                if sp in only_jp:
                    only_jp.remove(sp)
            shared_jp.extend(jp_alpha_s)
            shared_en.extend(en_alpha_s)
            only_jp.extend(jp_alpha_o)

            evidence_split_jp_alnum(jp_t, en_t, en_text, shared_jp, only_jp, shared_en)

            only_jp = filter_inner_only(only_jp, shared_jp)
            only_en = filter_inner_only(only_en, shared_en)

            # Apply highlighting: yellow for shared, red for mismatch
            apply_highlighting_to_cell(jp_cell, shared_jp, only_jp, default_color=WHITE)
            apply_highlighting_to_cell(en_cell, shared_en, only_en, default_color=WHITE)

            # Mark rows without mismatches for deletion
            if not only_jp and not only_en:
                rows_to_delete_idx.append(row_idx)

        # Delete clean rows from bottom to top
        for idx in reversed(rows_to_delete_idx):
            tbl._tbl.remove(tbl.rows[idx]._tr)

        apply_dark_table_style(tbl)
        distribute_column_widths(tbl)

    doc_mismatch.save(out_path_mismatch)


# ========== 11. CLI ==========

def main():
    path = input("DOCX path: ").strip().strip('"')
    if not os.path.isfile(path):
        print("File not found.")
        return

    # detect direction BEFORE processing (for filename tag)
    doc = Document(path)
    jp_col, en_col = detect_column_lang(doc, max_rows=50)
    tag = direction_tag(jp_col, en_col)

    base = os.path.splitext(path)[0]
    out_mismatch = base + f"_refsign_mismatch_{tag}.docx"
    out_all = base + f"_refsign_all_{tag}.docx"
    
    process_docx(path, out_mismatch, out_all)

    print(f"Detected TL direction: {tag}")
    print(f"Mismatch-only -> {out_mismatch}")
    print(f"All rows      -> {out_all}")

if __name__ == "__main__":
    main()
