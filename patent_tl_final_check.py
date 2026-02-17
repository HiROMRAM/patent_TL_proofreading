#!/usr/bin/env python3
"""
Patent Translation Final QA Check Script (Unified JP2EN / EN2JP)
Auto-detects translation direction from the input docx.

Input: docx file with 2-column table (Source | Target)
Output: xlsx file with 3 columns (Source | Target | Issues)

Processing:
- Non-claim segments: 7-row window, 5-row block, 2-row overlap
- Claim segments: grouped by 【請求項n】 markers, 3 claims per block
- Prompt caching for cost reduction

Usage:
  Interactive: python patent_tl_check.py  (then drag & drop file)
  CLI: python patent_tl_check.py input.docx [-o output.xlsx]
"""

import argparse
import json
import re
import sys
import unicodedata
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font

# ==========================================
# CONFIGURATION
# ==========================================
MODEL_NAME = "claude-opus-4-6"
MAX_TOKENS = 4096
WINDOW_SIZE = 7       # Context window for non-claim segments
BLOCK_SIZE = 5        # Block size for non-claim segments
CLAIMS_PER_BATCH = 3  # Claims per QA check batch

# Regex for claim markers
CLAIM_MARKER_RE = re.compile(r'【請求項[０-９\d]+】')


# ==========================================
# SYSTEM PROMPTS (direction-dependent)
# ==========================================
SYSTEM_PROMPT_JP2EN = """You are a patent translation QA checker for Japanese to English translations.
You check translation segments for CLEAR ERRORS ONLY.

Check ONLY for these types of clear errors:
1. 訳抜け (Omission): Information in JP that is missing in EN
2. 数字・単位の不一致 (Number/Unit mismatch): Numbers or units differ between JP and EN
3. 明らかな誤訳 (Clear mistranslation): Obviously wrong translation or established term not used
4. 明らかな文法エラー (Clear grammar error): Obvious grammatical mistakes in EN
5. 主語/目的語の取り違え (Subject/Object confusion): Subject or object incorrectly translated

IMPORTANT:
- Report clear, definite errors with full confidence.
- If you suspect an error but are not certain, report it with "(uncertain)" prepended to the detail field.
- Do NOT report mere stylistic preferences or minor improvements.
- This is error detection, NOT polishing.
- Only check the target rows specified, use other rows for context only.

Respond in JSON format:
{
  "issues": [
    {
      "row": <row_number>,
      "type": "<error_type>",
      "detail": "<brief description of the error and location>"
    }
  ]
}

If no clear errors are found in the target rows, respond with:
{"issues": []}
"""

SYSTEM_PROMPT_EN2JP = """You are a patent translation QA checker for English to Japanese translations.
You check translation segments for CLEAR ERRORS ONLY.

Each segment has:
- EN: the English source text
- JP: the Japanese translation to be checked

Check ONLY for these types of clear errors:
1. 訳抜け (Omission): Information in EN source that is missing in JP translation
2. 数字・単位の不一致 (Number/Unit mismatch): Numbers or units in JP differ from EN source
3. 明らかな誤訳 (Clear mistranslation): Obviously wrong JP translation, or established patent term not used
4. 明らかな文法エラー (Clear grammar error): Obvious grammatical mistakes in the JP translation (e.g. broken particles, wrong verb form)
5. 主語/目的語の取り違え (Subject/Object confusion): Subject or object incorrectly translated in JP

IMPORTANT:
- Report clear, definite errors with full confidence.
- If you suspect an error but are not certain, report it with "(uncertain)" prepended to the detail field.
- Do NOT report mere stylistic preferences or minor improvements.
- This is error detection, NOT polishing.
- Only check the target rows specified, use other rows for context only.
- Respond in Japanese for the detail field.

Respond in JSON format:
{
  "issues": [
    {
      "row": <row_number>,
      "type": "<error_type>",
      "detail": "<brief description of the error and location, in Japanese>"
    }
  ]
}

If no clear errors are found in the target rows, respond with:
{"issues": []}
"""


# ==========================================
# DIRECTION DETECTION (from ref_sign_checker)
# ==========================================

def _nfkc(s: str) -> str:
    return unicodedata.normalize("NFKC", s or "")


def _is_japanese_char(ch: str) -> bool:
    code = ord(ch)
    if 0x3040 <= code <= 0x309F:
        return True  # Hiragana
    if 0x30A0 <= code <= 0x30FF:
        return True  # Katakana
    if 0x4E00 <= code <= 0x9FFF:
        return True  # Kanji
    if 0xFF00 <= code <= 0xFFEF:
        return True  # Full-width forms
    return False


def _jp_score(text: str) -> float:
    chars = [c for c in text if not c.isspace()]
    if not chars:
        return 0.0
    return sum(1 for c in chars if _is_japanese_char(c)) / len(chars)


def _en_score(text: str) -> float:
    chars = [c for c in text if not c.isspace()]
    if not chars:
        return 0.0
    return sum(1 for c in chars if ("A" <= c <= "Z") or ("a" <= c <= "z")) / len(chars)


def detect_column_lang(doc: Document, max_rows: int = 50) -> tuple[int, int]:
    """Detect which column is JP and which is EN.

    Returns (jp_col, en_col) as 0-based indices.
    """
    for tbl in doc.tables:
        col_scores = [{"jp": 0.0, "en": 0.0}, {"jp": 0.0, "en": 0.0}]
        checked = 0

        for row in tbl.rows:
            if len(row.cells) < 2:
                continue
            t0 = _nfkc(row.cells[0].text)
            t1 = _nfkc(row.cells[1].text)

            col_scores[0]["jp"] += _jp_score(t0)
            col_scores[0]["en"] += _en_score(t0)
            col_scores[1]["jp"] += _jp_score(t1)
            col_scores[1]["en"] += _en_score(t1)

            checked += 1
            if checked >= max_rows:
                break

        if checked > 0:
            jp_col = 0 if col_scores[0]["jp"] >= col_scores[1]["jp"] else 1
            en_col = 1 - jp_col
            return jp_col, en_col

    return 0, 1  # fallback


def detect_direction(doc: Document, max_rows: int = 50) -> tuple[int, int, str]:
    """Detect translation direction from the docx.

    Returns (jp_col, en_col, direction) where direction is 'JP2EN' or 'EN2JP'.
    JP2EN = col0 is JP (source), col1 is EN (target)  → jp_col < en_col
    EN2JP = col0 is EN (source), col1 is JP (target)  → en_col < jp_col
    """
    jp_col, en_col = detect_column_lang(doc, max_rows)
    direction = "JP2EN" if jp_col < en_col else "EN2JP"
    return jp_col, en_col, direction


# ==========================================
# DOCX EXTRACTION
# ==========================================

def extract_table_from_docx(docx_path: str) -> list[tuple[str, str]]:
    """Extract text pairs from 2-column docx table.

    Returns list of (col0_text, col1_text) tuples.
    """
    doc = Document(docx_path)

    if not doc.tables:
        raise ValueError("No tables found in the document")

    table = doc.tables[0]
    pairs = []

    # Detect and skip header row
    first_cell = table.rows[0].cells[0].text.strip().lower()
    header_keywords = ('source', 'source segment', 'en', 'english', '原文', '英語',
                       'src', 'japanese', 'jp', 'ja', '日本語', 'target', 'target segment')
    start_row = 1 if first_cell in header_keywords else 0

    for i, row in enumerate(table.rows):
        if i < start_row:
            continue
        cells = row.cells
        if len(cells) >= 2:
            c0 = cells[0].text.strip()
            c1 = cells[1].text.strip()
            if c0 or c1:  # skip completely empty rows
                pairs.append((c0, c1))

    return pairs


# ==========================================
# CLAIM DETECTION & BATCHING
# ==========================================

def get_claim_number(text: str) -> int | None:
    """Extract claim number from text, or None if not a claim."""
    match = CLAIM_MARKER_RE.search(text)
    if match:
        raw = match.group()
        digits = raw.replace('【請求項', '').replace('】', '')
        digits = digits.translate(str.maketrans('０１２３４５６７８９', '0123456789'))
        return int(digits)
    return None


def group_into_check_batches(pairs: list[tuple[str, str]], jp_col: int) -> list[dict]:
    """Group rows into QA check batches, treating claims specially.

    Returns list of batch dicts:
    {
        'check_indices': [row_indices to check],    # 0-indexed
        'context_indices': [row_indices for context], # 0-indexed
        'is_claim': bool
    }
    """
    n = len(pairs)

    # Tag each row
    tagged = []  # (row_index, claim_number_or_None)
    for i, pair in enumerate(pairs):
        jp_text = pair[jp_col]
        claim_num = get_claim_number(jp_text)
        tagged.append((i, claim_num))

    # Identify claim zones: contiguous runs of claim/continuation rows
    claim_groups = []  # list of lists of row indices, one per claim
    current_claim_rows = []
    current_claim_num = None
    in_claim_zone = False

    for row_idx, claim_num in tagged:
        if claim_num is not None:
            in_claim_zone = True
            if claim_num != current_claim_num:
                if current_claim_rows:
                    claim_groups.append(current_claim_rows)
                current_claim_rows = [row_idx]
                current_claim_num = claim_num
            else:
                current_claim_rows.append(row_idx)
        elif in_claim_zone:
            # Check for section boundary
            jp_text = pairs[row_idx][jp_col]
            section_marker = re.search(r'【[^請][^求][^項].*?】', jp_text)
            if section_marker:
                if current_claim_rows:
                    claim_groups.append(current_claim_rows)
                    current_claim_rows = []
                    current_claim_num = None
                in_claim_zone = False
            else:
                current_claim_rows.append(row_idx)
        # non-claim rows outside claim zone handled below

    if current_claim_rows:
        claim_groups.append(current_claim_rows)

    # Build set of all claim row indices
    claim_row_set = set()
    for group in claim_groups:
        claim_row_set.update(group)

    # Build batches
    batches = []

    # Claim batches: CLAIMS_PER_BATCH claims at a time
    for i in range(0, len(claim_groups), CLAIMS_PER_BATCH):
        chunk = claim_groups[i:i + CLAIMS_PER_BATCH]
        check_indices = []
        for group in chunk:
            check_indices.extend(group)

        # Context: 1 row before first, 1 row after last
        first_idx = min(check_indices)
        last_idx = max(check_indices)
        context_indices = list(check_indices)
        if first_idx > 0:
            context_indices.insert(0, first_idx - 1)
        if last_idx < n - 1:
            context_indices.append(last_idx + 1)

        batches.append({
            'check_indices': check_indices,
            'context_indices': sorted(set(context_indices)),
            'is_claim': True
        })

    # Non-claim batches: sliding window with overlap
    non_claim_indices = [i for i in range(n) if i not in claim_row_set]

    block_start = 0
    while block_start < len(non_claim_indices):
        block_end = min(block_start + BLOCK_SIZE, len(non_claim_indices))
        block_rows = non_claim_indices[block_start:block_end]

        # Window: extend for context
        overlap = WINDOW_SIZE - BLOCK_SIZE
        context_before = overlap // 2
        context_after = overlap - context_before

        first_idx = block_rows[0]
        last_idx = block_rows[-1]
        window_start = max(0, first_idx - context_before)
        window_end = min(n - 1, last_idx + context_after)
        context_indices = list(range(window_start, window_end + 1))

        batches.append({
            'check_indices': block_rows,
            'context_indices': sorted(set(context_indices)),
            'is_claim': False
        })

        block_start = block_end

    return batches


# ==========================================
# API INTERACTION
# ==========================================

def build_user_message(
    pairs: list[tuple[str, str]],
    batch: dict,
    jp_col: int,
    en_col: int,
    direction: str,
) -> str:
    """Build the user message for a QA check batch.

    Source language is listed first in each row for clarity.
    """
    context_rows = batch['context_indices']
    check_rows = batch['check_indices']

    if direction == "JP2EN":
        # Source=JP, Target=EN
        context_text = "\n".join(
            f"[Row {i + 1}]\nJP: {pairs[i][jp_col]}\nEN: {pairs[i][en_col]}"
            for i in context_rows
        )
    else:
        # Source=EN, Target=JP
        context_text = "\n".join(
            f"[Row {i + 1}]\nEN: {pairs[i][en_col]}\nJP: {pairs[i][jp_col]}"
            for i in context_rows
        )

    check_rows_str = ", ".join(str(i + 1) for i in check_rows)

    return f"""Check the following translation segments.
Target rows to check: {check_rows_str}
(Other rows are provided for context only)

<segments>
{context_text}
</segments>"""


def check_batch(
    client: anthropic.Anthropic,
    system_prompt_blocks: list[dict],
    pairs: list[tuple[str, str]],
    batch: dict,
    jp_col: int,
    en_col: int,
    direction: str,
) -> list[dict]:
    """Send a batch to Claude API for QA checking."""
    user_message = build_user_message(pairs, batch, jp_col, en_col, direction)

    try:
        response = client.messages.create(
            model=MODEL_NAME,
            max_tokens=MAX_TOKENS,
            system=system_prompt_blocks,
            messages=[{"role": "user", "content": user_message}]
        )

        # Log cache performance
        usage = response.usage
        cache_read = getattr(usage, 'cache_read_input_tokens', 0)
        cache_create = getattr(usage, 'cache_creation_input_tokens', 0)
        if cache_create > 0:
            print(f" [cache: wrote {cache_create}tok]", end="")
        elif cache_read > 0:
            print(f" [cache: read {cache_read}tok]", end="")

        response_text = response.content[0].text

        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            try:
                result = json.loads(json_match.group())
                return result.get("issues", [])
            except json.JSONDecodeError:
                print(f" [Warning: failed to parse JSON]", end="")
                return []
        return []

    except Exception as e:
        print(f"\n  [Error] {e}")
        raise


# ==========================================
# MAIN PROCESSING
# ==========================================

def process_document(
    docx_path: str,
    output_path: str,
) -> None:
    """Process document and generate QA report."""

    # Extract pairs
    print(f"Reading {docx_path}...")
    pairs = extract_table_from_docx(docx_path)
    total_rows = len(pairs)
    print(f"Found {total_rows} translation pairs")

    # Detect direction from the docx (using character-ratio scoring)
    doc = Document(docx_path)
    jp_col, en_col, direction = detect_direction(doc, max_rows=50)

    if direction == "JP2EN":
        print(f"  Detected direction: JP → EN  (col0=JP, col1=EN)")
    else:
        print(f"  Detected direction: EN → JP  (col0=EN, col1=JP)")

    # Select system prompt based on direction
    system_prompt_text = SYSTEM_PROMPT_JP2EN if direction == "JP2EN" else SYSTEM_PROMPT_EN2JP

    # Initialize API client
    client = anthropic.Anthropic()

    # Build system prompt with caching
    system_prompt_blocks = [
        {
            "type": "text",
            "text": system_prompt_text,
            "cache_control": {"type": "ephemeral"}
        }
    ]

    # Group into batches
    batches = group_into_check_batches(pairs, jp_col)

    claim_batches = sum(1 for b in batches if b['is_claim'])
    non_claim_batches = len(batches) - claim_batches
    claim_count = sum(
        1 for pair in pairs if get_claim_number(pair[jp_col]) is not None
    )
    print(f"  Claims detected: {claim_count}")
    print(f"  Batches: {len(batches)} ({claim_batches} claim, {non_claim_batches} non-claim)")
    print(f"  Model: {MODEL_NAME}")

    # Process batches
    all_issues: dict[int, list[str]] = {}  # 1-indexed row -> list of issue strings
    print("\nProcessing...\n")

    for batch_num, batch in enumerate(batches, 1):
        check_indices = batch['check_indices']

        if batch['is_claim']:
            claim_ids = [
                CLAIM_MARKER_RE.search(pairs[i][jp_col]).group()
                for i in check_indices
                if CLAIM_MARKER_RE.search(pairs[i][jp_col])
            ]
            label = f"Claims {', '.join(claim_ids)}" if claim_ids else f"{len(check_indices)} rows"
        else:
            label = f"rows {check_indices[0]+1}-{check_indices[-1]+1}"

        print(f"  Batch {batch_num}/{len(batches)} [{label}]...", end="", flush=True)

        issues = check_batch(client, system_prompt_blocks, pairs, batch,
                             jp_col, en_col, direction)

        # Collect issues (Claude returns 1-indexed row numbers)
        valid_check_rows = set(i + 1 for i in check_indices)
        for issue in issues:
            row_num = issue.get("row")
            if row_num and row_num in valid_check_rows:
                issue_text = f"[{issue.get('type', 'Error')}] {issue.get('detail', '')}"
                if row_num not in all_issues:
                    all_issues[row_num] = []
                if issue_text not in all_issues[row_num]:
                    all_issues[row_num].append(issue_text)

        issue_count = len(issues)
        if issue_count > 0:
            print(f" {issue_count} issue(s) found")
        else:
            print(" OK")

    # Create output Excel (two sheets: All, IssuesOnly)
    print(f"\nCreating output: {output_path}")
    wb = Workbook()

    # Headers & column mapping depend on direction (source first)
    if direction == "JP2EN":
        headers = ['Japanese (Source)', 'English (Translation)', 'Issues']
        src_col_idx, tgt_col_idx = jp_col, en_col
    else:
        headers = ['English (Source)', 'Japanese (Translation)', 'Issues']
        src_col_idx, tgt_col_idx = en_col, jp_col

    header_font = Font(bold=True)
    wrap_top = Alignment(wrap_text=True, vertical='top')

    def _is_uncertain(issue_text: str) -> bool:
        """Check if a single issue string is an (uncertain) flag."""
        # Match pattern: [ErrorType] (uncertain) ...
        return "(uncertain)" in issue_text

    def _has_definite_issue(row_id: int) -> bool:
        """True if the row has at least one non-uncertain issue."""
        if row_id not in all_issues:
            return False
        return any(not _is_uncertain(t) for t in all_issues[row_id])

    # --- Sheet 1: All rows (including uncertain) ---
    ws_all = wb.active
    ws_all.title = "All"

    for col_idx, header in enumerate(headers, 1):
        cell = ws_all.cell(row=1, column=col_idx, value=header)
        cell.font = header_font

    for i, pair in enumerate(pairs):
        row_num = i + 2
        row_id = i + 1
        ws_all.cell(row=row_num, column=1, value=pair[src_col_idx])
        ws_all.cell(row=row_num, column=2, value=pair[tgt_col_idx])
        if row_id in all_issues:
            ws_all.cell(row=row_num, column=3, value="\n".join(all_issues[row_id]))

    ws_all.column_dimensions['A'].width = 60
    ws_all.column_dimensions['B'].width = 60
    ws_all.column_dimensions['C'].width = 50
    for row in ws_all.iter_rows(min_row=1, max_row=total_rows + 1):
        for cell in row:
            cell.alignment = wrap_top

    # --- Sheet 2: IssuesOnly (rows with at least one definite issue) ---
    ws_issues = wb.create_sheet("IssuesOnly")

    for col_idx, header in enumerate(headers, 1):
        cell = ws_issues.cell(row=1, column=col_idx, value=header)
        cell.font = header_font

    out_row = 2
    definite_count = 0
    for i, pair in enumerate(pairs):
        row_id = i + 1
        if not _has_definite_issue(row_id):
            continue
        ws_issues.cell(row=out_row, column=1, value=pair[src_col_idx])
        ws_issues.cell(row=out_row, column=2, value=pair[tgt_col_idx])
        ws_issues.cell(row=out_row, column=3, value="\n".join(all_issues[row_id]))
        out_row += 1
        definite_count += 1

    ws_issues.column_dimensions['A'].width = 60
    ws_issues.column_dimensions['B'].width = 60
    ws_issues.column_dimensions['C'].width = 50
    for row in ws_issues.iter_rows(min_row=1, max_row=out_row):
        for cell in row:
            cell.alignment = wrap_top

    # Save
    try:
        wb.save(output_path)
        print(f"Saved to: {output_path}")
    except PermissionError:
        stem = Path(output_path).stem
        suffix = Path(output_path).suffix
        parent = Path(output_path).parent
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        fallback = parent / f"{stem}_{ts}{suffix}"
        wb.save(str(fallback))
        print(f"  [!] Could not save to {output_path} (file open?)")
        print(f"  Saved to: {fallback}")

    # Summary
    total_issues = sum(len(v) for v in all_issues.values())
    uncertain_issues = sum(
        sum(1 for t in v if _is_uncertain(t)) for v in all_issues.values()
    )
    definite_issues = total_issues - uncertain_issues
    rows_with_any = len(all_issues)
    print(f"\nDone! {total_issues} issues in {rows_with_any} rows "
          f"({definite_issues} definite, {uncertain_issues} uncertain). "
          f"IssuesOnly sheet: {definite_count} rows.")


# ==========================================
# CLI & INTERACTIVE ENTRY POINTS
# ==========================================

def main_cli():
    """CLI mode with arguments."""
    parser = argparse.ArgumentParser(
        description="Patent Translation QA Check using Claude API (auto-detects JP2EN / EN2JP)"
    )
    parser.add_argument("input", help="Input docx file with translation table")
    parser.add_argument("-o", "--output", help="Output xlsx file")

    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    output_path = args.output or str(
        input_path.parent / (input_path.stem + "_qa_result.xlsx")
    )

    process_document(str(input_path), str(output_path))


def main_interactive():
    """Interactive mode."""
    print("=" * 60)
    print("  Patent Translation QA Check (auto-detect direction)")
    print("=" * 60)
    print(f"  Model: {MODEL_NAME}")
    print("=" * 60)
    print()

    input_str = input("[1] Enter docx file path (or drag & drop):\n    > ").strip()
    input_str = input_str.strip('"').strip("'")

    input_path = Path(input_str)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        input("\nPress Enter to exit...")
        sys.exit(1)

    if input_path.suffix.lower() != '.docx':
        print(f"Error: File must be .docx format")
        input("\nPress Enter to exit...")
        sys.exit(1)

    output_path = input_path.parent / (input_path.stem + "_qa_result.xlsx")

    print()
    process_document(str(input_path), str(output_path))

    print()
    input("Press Enter to exit...")


if __name__ == "__main__":
    if len(sys.argv) > 1:
        main_cli()
    else:
        main_interactive()
